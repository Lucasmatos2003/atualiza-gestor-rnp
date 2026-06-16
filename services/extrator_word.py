import re
import unicodedata
import pandas as pd
from docx import Document


COLUNAS_PADRAO = [
    "inep",
    "escola",
    "gestor",
    "telefone",
    "telefone_2",
    "telefone_3",
    "outros_telefones",
    "email",
    "municipio",
    "origem",
    "fonte_contato"
]


def normalizar_texto(texto):
    if texto is None or pd.isna(texto):
        return ""

    texto = str(texto).strip()
    texto = unicodedata.normalize("NFD", texto)
    texto = texto.encode("ascii", "ignore")
    texto = texto.decode("utf-8")
    texto = re.sub(r"\s+", " ", texto)

    return texto.strip()


def normalizar_minusculo(texto):
    return normalizar_texto(texto).lower()


def limpar_valor(valor):
    if valor is None or pd.isna(valor):
        return None

    valor = str(valor).replace("\xa0", " ").strip()
    valor = re.sub(r"\s+", " ", valor)

    if valor.lower() in [
        "",
        "nan",
        "none",
        "-",
        "não possui",
        "nao possui",
        "não informado",
        "nao informado"
    ]:
        return None

    return valor


def limpar_inep(valor):
    if valor is None or pd.isna(valor):
        return None

    valor = str(valor)

    match = re.search(
        r"INEP\s*[:\-]?\s*(\d{7,8})",
        valor,
        flags=re.IGNORECASE
    )

    if match:
        return match.group(1)

    match = re.search(r"\b\d{7,8}\b", valor)

    if match:
        return match.group(0)

    valor = re.sub(r"\D", "", valor)

    if len(valor) >= 7:
        return valor[:8]

    return None


def extrair_todos_emails(valor):
    if valor is None or pd.isna(valor):
        return []

    valor = str(valor)
    valor = valor.replace("\n", " ")
    valor = re.sub(r"\s*@\s*", "@", valor)
    valor = re.sub(r"\s*\.\s*", ".", valor)

    emails = re.findall(
        r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
        valor
    )

    emails_limpos = []

    for email in emails:
        email = email.lower()

        if email not in emails_limpos:
            emails_limpos.append(email)

    return emails_limpos


def limpar_email(valor):
    emails = extrair_todos_emails(valor)

    if emails:
        return emails[0]

    return None


def limpar_telefone(valor):
    if valor is None or pd.isna(valor):
        return None

    valor = str(valor)

    telefones = re.findall(
        r"(?:\+?55\s*)?(?:\(?\d{2}\)?\s*)?\d?\s?\d{4,5}[-\s]?\d{4}",
        valor
    )

    candidatos = []

    for telefone in telefones:
        numero = re.sub(r"\D", "", telefone)

        if numero.startswith("55") and len(numero) in [12, 13]:
            numero = numero[2:]

        if numero.startswith("0") and len(numero) in [11, 12]:
            numero = numero[1:]

        if len(numero) in [8, 9, 10, 11]:
            candidatos.append(numero)

    if candidatos:
        return candidatos[0]

    numero = re.sub(r"\D", "", valor)

    if numero.startswith("55") and len(numero) in [12, 13]:
        numero = numero[2:]

    if numero.startswith("0") and len(numero) in [11, 12]:
        numero = numero[1:]

    if len(numero) in [8, 9, 10, 11]:
        return numero

    return None


def extrair_todos_telefones(valor):
    if valor is None or pd.isna(valor):
        return []

    valor = str(valor)

    encontrados = re.findall(
        r"(?:\+?55\s*)?(?:\(?\d{2}\)?\s*)?\d?\s?\d{4,5}[-\s]?\d{4}",
        valor
    )

    telefones = []

    for item in encontrados:
        telefone = limpar_telefone(item)

        if telefone and telefone not in telefones:
            telefones.append(telefone)

    if not telefones:
        telefone = limpar_telefone(valor)

        if telefone and telefone not in telefones:
            telefones.append(telefone)

    return telefones


def ordenar_telefones(telefones):
    telefones_unicos = []

    for telefone in telefones:
        if not telefone:
            continue

        telefone = re.sub(r"\D", "", str(telefone))

        if telefone.startswith("55") and len(telefone) in [12, 13]:
            telefone = telefone[2:]

        if telefone.startswith("0") and len(telefone) in [11, 12]:
            telefone = telefone[1:]

        if len(telefone) < 8:
            continue

        if telefone not in telefones_unicos:
            telefones_unicos.append(telefone)

    def prioridade(telefone):
        if len(telefone) == 11 and telefone[2] == "9":
            return 4

        if len(telefone) == 11:
            return 3

        if len(telefone) == 10:
            return 2

        if len(telefone) == 9:
            return 1

        return 0

    return sorted(
        telefones_unicos,
        key=prioridade,
        reverse=True
    )


def possui_email(texto):
    if not texto:
        return False

    return bool(extrair_todos_emails(texto))


def possui_telefone(texto):
    if not texto:
        return False

    return bool(extrair_todos_telefones(texto))


def limpar_nome_gestor(valor):
    if valor is None or pd.isna(valor):
        return None

    texto = str(valor).replace("\n", " ").strip()

    texto = re.sub(
        r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
        " ",
        texto
    )

    texto = re.sub(
        r"(?:\+?55\s*)?(?:\(?\d{2}\)?\s*)?\d?\s?\d{4,5}[-\s]?\d{4}",
        " ",
        texto
    )

    texto = re.sub(r"\s+", " ", texto).strip(" -()")

    if texto.lower() in [
        "",
        "nan",
        "none",
        "-",
        "não informado",
        "nao informado",
        "não possui",
        "nao possui"
    ]:
        return None

    return texto


def montar_dataframe(dados):
    if not dados:
        return pd.DataFrame(columns=COLUNAS_PADRAO)

    df = pd.DataFrame(dados)

    for coluna in COLUNAS_PADRAO:
        if coluna not in df.columns:
            df[coluna] = None

    df = df[COLUNAS_PADRAO]

    df = df.drop_duplicates(
        subset=[
            "inep",
            "escola",
            "gestor",
            "telefone",
            "email"
        ],
        keep="first"
    )

    df = df.reset_index(drop=True)

    return df


def texto_celula(celula):
    if celula is None:
        return ""

    texto = celula.text.replace("\xa0", " ").strip()
    texto = re.sub(r"[ \t]+", " ", texto)

    linhas = []

    for linha in texto.splitlines():
        linha = linha.strip()

        if linha:
            linhas.append(linha)

    return "\n".join(linhas)


def obter_texto_documento(documento):
    partes = []

    for paragrafo in documento.paragraphs:
        texto = paragrafo.text.strip()

        if texto:
            partes.append(texto)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                texto = celula.text.strip()

                if texto:
                    partes.append(texto)

    return "\n".join(partes)


def primeira_linha_valida(texto):
    if not texto:
        return None

    for linha in str(texto).splitlines():
        linha = limpar_valor(linha)

        if linha:
            return linha

    return None


def extrair_nome_escola_de_celula_unidade(texto):
    if not texto:
        return None

    linhas = [
        linha.strip()
        for linha in str(texto).splitlines()
        if linha.strip()
    ]

    if not linhas:
        return None

    primeira = linhas[0]

    primeira = re.sub(r"^\d+\s*[-–]\s*", "", primeira).strip()

    if "INEP" in primeira.upper():
        primeira = re.split(
            r"\bINEP\b",
            primeira,
            flags=re.IGNORECASE
        )[0].strip()

    escola = primeira

    if len(linhas) > 1 and "INEP" not in linhas[1].upper():
        segunda = linhas[1].strip()

        if not re.search(r"\bSEC\b", segunda, flags=re.IGNORECASE):
            if not re.search(r"ENSINO|EJA|EDUCAÇÃO|EDUCACAO", segunda, flags=re.IGNORECASE):
                escola = f"{escola} {segunda}"

    escola = re.sub(r"\s+", " ", escola).strip()

    return limpar_valor(escola)


def detectar_juazeiro_estadual(documento):
    texto = obter_texto_documento(documento)
    texto_norm = normalizar_minusculo(texto)

    condicoes = [
        "lista unidades escolares sede e territorio" in texto_norm
        or "lista unidades escolares sede e território" in texto_norm,
        "juazeiro" in texto_norm,
        "unidade escolar" in texto_norm,
        "direcao" in texto_norm or "direção" in texto_norm,
        "telefone" in texto_norm,
        "e-mail" in texto_norm or "email" in texto_norm
    ]

    return all(condicoes)


def encontrar_indice_unidade(celulas_texto):
    for i, texto in enumerate(celulas_texto):
        texto_norm = normalizar_minusculo(texto)

        if "inep" not in texto_norm:
            continue

        if "unidade escolar" in texto_norm:
            continue

        if any(palavra in texto_norm for palavra in [
            "colegio",
            "colégio",
            "centro",
            "complexo",
            "ceep",
            "cetep"
        ]):
            return i

    for i, texto in enumerate(celulas_texto):
        texto_norm = normalizar_minusculo(texto)

        if "inep" in texto_norm and "unidade escolar" not in texto_norm:
            return i

    return None


def encontrar_indice_email(celulas_texto):
    for i, texto in enumerate(celulas_texto):
        if possui_email(texto):
            return i

    return None


def encontrar_texto_gestor(celulas_texto, indice_unidade, indice_email):
    candidatos = []

    for i, texto in enumerate(celulas_texto):
        if not texto:
            continue

        if i == indice_unidade:
            continue

        if i == indice_email:
            continue

        texto_norm = normalizar_minusculo(texto)

        if "inep" in texto_norm:
            continue

        if "sec" in texto_norm and len(texto_norm) < 20:
            continue

        if "unidade escolar" in texto_norm:
            continue

        if "direcao" in texto_norm or "direção" in texto_norm:
            continue

        if "telefone" in texto_norm:
            continue

        if "e-mail" in texto_norm or "email" in texto_norm:
            continue

        if "anexo" in texto_norm:
            continue

        if possui_email(texto):
            continue

        if possui_telefone(texto):
            continue

        primeira = primeira_linha_valida(texto)

        if not primeira:
            continue

        primeira_norm = normalizar_minusculo(primeira)

        if any(palavra in primeira_norm for palavra in [
            "colegio",
            "colégio",
            "ensino",
            "educacao",
            "educação",
            "eja",
            "sec",
            "inep"
        ]):
            continue

        candidatos.append(primeira)

    if candidatos:
        return candidatos[0]

    return None


def extrair_juazeiro_estadual(documento):
    dados = []

    for tabela in documento.tables:
        for linha in tabela.rows:
            celulas_texto = [
                texto_celula(celula)
                for celula in linha.cells
            ]

            if not celulas_texto:
                continue

            indice_unidade = encontrar_indice_unidade(celulas_texto)

            if indice_unidade is None:
                continue

            unidade = celulas_texto[indice_unidade]

            escola = extrair_nome_escola_de_celula_unidade(unidade)
            inep = limpar_inep(unidade)

            if not escola or not inep:
                continue

            indice_email = encontrar_indice_email(celulas_texto)

            email_texto = celulas_texto[indice_email] if indice_email is not None else ""
            email_principal = limpar_email(email_texto)

            gestor_texto = encontrar_texto_gestor(
                celulas_texto=celulas_texto,
                indice_unidade=indice_unidade,
                indice_email=indice_email
            )

            gestor = limpar_nome_gestor(gestor_texto)

            telefones = []

            for i, texto in enumerate(celulas_texto):
                if i == indice_unidade:
                    continue

                if i == indice_email:
                    continue

                telefones.extend(
                    extrair_todos_telefones(texto)
                )

            telefones = ordenar_telefones(telefones)

            telefone_1 = telefones[0] if len(telefones) >= 1 else None
            telefone_2 = telefones[1] if len(telefones) >= 2 else None
            telefone_3 = telefones[2] if len(telefones) >= 3 else None
            outros_telefones = "; ".join(telefones[3:]) if len(telefones) > 3 else None

            dados.append({
                "inep": inep,
                "escola": escola,
                "gestor": gestor,
                "telefone": telefone_1,
                "telefone_2": telefone_2,
                "telefone_3": telefone_3,
                "outros_telefones": outros_telefones,
                "email": email_principal,
                "municipio": "Juazeiro",
                "origem": "Juazeiro Estadual DOCX",
                "fonte_contato": "Juazeiro Estadual DOCX"
            })

    return montar_dataframe(dados)


def extrair_word_generico(documento):
    dados = []

    for tabela in documento.tables:
        if not tabela.rows:
            continue

        cabecalho = [
            normalizar_minusculo(celula.text)
            for celula in tabela.rows[0].cells
        ]

        texto_cabecalho = " ".join(cabecalho)

        if not (
            "escola" in texto_cabecalho
            or "unidade" in texto_cabecalho
        ):
            continue

        idx_inep = None
        idx_escola = None
        idx_gestor = None
        idx_telefone = None
        idx_email = None
        idx_municipio = None

        for i, coluna in enumerate(cabecalho):
            if "inep" in coluna:
                idx_inep = i

            if "escola" in coluna or "unidade" in coluna:
                idx_escola = i

            if (
                "gestor" in coluna
                or "diretor" in coluna
                or "direcao" in coluna
                or "direção" in coluna
            ):
                idx_gestor = i

            if "telefone" in coluna or "contato" in coluna or "celular" in coluna:
                idx_telefone = i

            if "email" in coluna or "e-mail" in coluna:
                idx_email = i

            if "municipio" in coluna or "município" in coluna or "cidade" in coluna:
                idx_municipio = i

        for linha in tabela.rows[1:]:
            celulas = linha.cells

            def valor(indice):
                if indice is None:
                    return None

                if indice >= len(celulas):
                    return None

                return texto_celula(celulas[indice])

            escola = limpar_valor(valor(idx_escola))
            gestor = limpar_nome_gestor(valor(idx_gestor))
            inep = limpar_inep(valor(idx_inep))
            email = limpar_email(valor(idx_email))
            municipio = limpar_valor(valor(idx_municipio))

            telefones = ordenar_telefones(
                extrair_todos_telefones(valor(idx_telefone))
            )

            telefone_1 = telefones[0] if len(telefones) >= 1 else None
            telefone_2 = telefones[1] if len(telefones) >= 2 else None
            telefone_3 = telefones[2] if len(telefones) >= 3 else None
            outros_telefones = "; ".join(telefones[3:]) if len(telefones) > 3 else None

            if not escola:
                continue

            dados.append({
                "inep": inep,
                "escola": escola,
                "gestor": gestor,
                "telefone": telefone_1,
                "telefone_2": telefone_2,
                "telefone_3": telefone_3,
                "outros_telefones": outros_telefones,
                "email": email,
                "municipio": municipio,
                "origem": "Word genérico",
                "fonte_contato": "Word genérico"
            })

    return montar_dataframe(dados)


def extrair_word(arquivo):
    arquivo.seek(0)

    documento = Document(arquivo)

    if detectar_juazeiro_estadual(documento):
        df_juazeiro = extrair_juazeiro_estadual(documento)

        if not df_juazeiro.empty:
            return df_juazeiro

    df_generico = extrair_word_generico(documento)

    if not df_generico.empty:
        return df_generico

    return pd.DataFrame(columns=COLUNAS_PADRAO)